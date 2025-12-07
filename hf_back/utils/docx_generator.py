"""
DOCX Generator Utility
Handles creation of Word documents (.docx) for all template types
"""

import io
import sys
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


class DocxGenerator:
    """Generate professional Word documents"""

    def __init__(self):
        """Initialize DOCX generator"""
        self.default_font = "Calibri"
        self.heading_font = "Arial"

    def _set_cell_border(self, cell, **kwargs):
        """
        Set cell border

        Args:
            cell: Table cell
            kwargs: Border properties (top, bottom, left, right)
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("left", "top", "right", "bottom"):
            if edge in kwargs:
                edge_data = kwargs.get(edge)
                edge_el = OxmlElement(f"w:{edge}")
                edge_el.set(qn("w:val"), "single")
                edge_el.set(qn("w:sz"), "4")
                edge_el.set(qn("w:space"), "0")
                edge_el.set(qn("w:color"), edge_data.get("color", "000000"))
                tcBorders.append(edge_el)

        tcPr.append(tcBorders)

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """Add styled heading to document"""
        heading = doc.add_heading(text, level=level)
        heading.style.font.name = self.heading_font
        heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        return heading

    def _add_paragraph(
        self,
        doc: Document,
        text: str,
        bold: bool = False,
        italic: bool = False,
        size: int = 11,
    ):
        """Add styled paragraph to document"""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = self.default_font
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        return para

    def generate_resume(self, data: Dict[str, Any]) -> io.BytesIO:
        """
        Generate resume document

        Args:
            data: Resume data containing:
                - personal_info: name, email, phone, location, linkedin, website
                - summary: Professional summary
                - experience: List of work experiences
                - education: List of education entries
                - skills: List of skills
                - certifications: List of certifications (optional)
                - projects: List of projects (optional)

        Returns:
            BytesIO buffer containing the .docx file
        """
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Personal Information (Header)
        personal = data.get("personal_info", {})

        # Name (Large and centered)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(personal.get("name", "Your Name"))
        name_run.font.name = self.heading_font
        name_run.font.size = Pt(24)
        name_run.bold = True
        name_run.font.color.rgb = RGBColor(0, 51, 102)

        # Contact Info (Centered)
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = []
        if personal.get("email"):
            contact_parts.append(personal["email"])
        if personal.get("phone"):
            contact_parts.append(personal["phone"])
        if personal.get("location"):
            contact_parts.append(personal["location"])

        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.name = self.default_font
        contact_run.font.size = Pt(10)

        # Links (Centered)
        if personal.get("linkedin") or personal.get("website"):
            links_para = doc.add_paragraph()
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            links_parts = []
            if personal.get("linkedin"):
                links_parts.append(f"LinkedIn: {personal['linkedin']}")
            if personal.get("website"):
                links_parts.append(f"Portfolio: {personal['website']}")

            links_run = links_para.add_run(" | ".join(links_parts))
            links_run.font.name = self.default_font
            links_run.font.size = Pt(10)
            links_run.font.color.rgb = RGBColor(0, 102, 204)

        doc.add_paragraph()  # Spacing

        # Professional Summary
        if data.get("summary"):
            self._add_heading(doc, "PROFESSIONAL SUMMARY", level=1)
            self._add_paragraph(doc, data["summary"])
            doc.add_paragraph()  # Spacing

        # Work Experience
        if data.get("experience"):
            self._add_heading(doc, "WORK EXPERIENCE", level=1)

            for exp in data["experience"]:
                # Job Title and Company
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(exp.get("title", "Position"))
                title_run.font.name = self.default_font
                title_run.font.size = Pt(12)
                title_run.bold = True

                # Company and Dates
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(
                    f"{exp.get('company', 'Company')} | "
                    f"{exp.get('start_date', 'Start')} - {exp.get('end_date', 'End')}"
                )
                company_run.font.name = self.default_font
                company_run.font.size = Pt(11)
                company_run.italic = True

                # Location
                if exp.get("location"):
                    location_run = company_para.add_run(f" | {exp['location']}")
                    location_run.font.name = self.default_font
                    location_run.font.size = Pt(11)
                    location_run.italic = True

                # Responsibilities/Achievements
                if exp.get("responsibilities"):
                    for resp in exp["responsibilities"]:
                        bullet_para = doc.add_paragraph(resp, style="List Bullet")
                        bullet_para.paragraph_format.left_indent = Inches(0.25)

                doc.add_paragraph()  # Spacing between jobs

        # Education
        if data.get("education"):
            self._add_heading(doc, "EDUCATION", level=1)

            for edu in data["education"]:
                # Degree
                degree_para = doc.add_paragraph()
                degree_run = degree_para.add_run(
                    f"{edu.get('degree', 'Degree')} in {edu.get('field', 'Field')}"
                )
                degree_run.font.name = self.default_font
                degree_run.font.size = Pt(12)
                degree_run.bold = True

                # School and Dates
                school_para = doc.add_paragraph()
                school_run = school_para.add_run(
                    f"{edu.get('school', 'School')} | "
                    f"{edu.get('graduation_date', 'Graduation Date')}"
                )
                school_run.font.name = self.default_font
                school_run.font.size = Pt(11)
                school_run.italic = True

                # GPA or Honors
                if edu.get("gpa") or edu.get("honors"):
                    details = []
                    if edu.get("gpa"):
                        details.append(f"GPA: {edu['gpa']}")
                    if edu.get("honors"):
                        details.append(edu["honors"])

                    details_para = doc.add_paragraph(" | ".join(details))
                    details_para.paragraph_format.left_indent = Inches(0.25)

                doc.add_paragraph()  # Spacing

        # Skills
        if data.get("skills"):
            self._add_heading(doc, "SKILLS", level=1)

            # Group skills by category if provided
            if isinstance(data["skills"], dict):
                for category, skills_list in data["skills"].items():
                    skills_para = doc.add_paragraph()
                    category_run = skills_para.add_run(f"{category}: ")
                    category_run.font.name = self.default_font
                    category_run.font.size = Pt(11)
                    category_run.bold = True

                    skills_run = skills_para.add_run(", ".join(skills_list))
                    skills_run.font.name = self.default_font
                    skills_run.font.size = Pt(11)
            else:
                # Simple list of skills
                skills_para = doc.add_paragraph(", ".join(data["skills"]))
                skills_para.paragraph_format.left_indent = Inches(0.25)

            doc.add_paragraph()  # Spacing

        # Certifications
        if data.get("certifications"):
            self._add_heading(doc, "CERTIFICATIONS", level=1)

            for cert in data["certifications"]:
                cert_para = doc.add_paragraph(style="List Bullet")
                cert_run = cert_para.add_run(
                    f"{cert.get('name', 'Certification')} - "
                    f"{cert.get('issuer', 'Issuer')}"
                )
                cert_run.font.name = self.default_font
                cert_run.font.size = Pt(11)

                if cert.get("date"):
                    date_run = cert_para.add_run(f" ({cert['date']})")
                    date_run.font.name = self.default_font
                    date_run.font.size = Pt(10)
                    date_run.italic = True

        # Projects
        if data.get("projects"):
            self._add_heading(doc, "PROJECTS", level=1)

            for proj in data["projects"]:
                # Project Title
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(proj.get("name", "Project"))
                proj_run.font.name = self.default_font
                proj_run.font.size = Pt(12)
                proj_run.bold = True

                # Description
                if proj.get("description"):
                    desc_para = doc.add_paragraph(proj["description"])
                    desc_para.paragraph_format.left_indent = Inches(0.25)

                # Technologies
                if proj.get("technologies"):
                    tech_para = doc.add_paragraph()
                    tech_para.paragraph_format.left_indent = Inches(0.25)
                    tech_label = tech_para.add_run("Technologies: ")
                    tech_label.font.size = Pt(10)
                    tech_label.italic = True
                    tech_list = tech_para.add_run(", ".join(proj["technologies"]))
                    tech_list.font.size = Pt(10)

                doc.add_paragraph()  # Spacing

        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def generate_cover_letter(self, data: Dict[str, Any]) -> io.BytesIO:
        """
        Generate cover letter document

        Args:
            data: Cover letter data containing:
                - name: Applicant name
                - address: Applicant address
                - email: Email
                - phone: Phone
                - date: Letter date
                - company: Company name
                - hiring_manager: Hiring manager name (optional)
                - position: Position applied for
                - content: Letter content (paragraphs)

        Returns:
            BytesIO buffer containing the .docx file
        """
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Applicant Info
        self._add_paragraph(doc, data.get("name", "Your Name"), bold=True, size=12)
        if data.get("address"):
            self._add_paragraph(doc, data["address"], size=10)

        contact_line = []
        if data.get("email"):
            contact_line.append(data["email"])
        if data.get("phone"):
            contact_line.append(data["phone"])
        if contact_line:
            self._add_paragraph(doc, " | ".join(contact_line), size=10)

        doc.add_paragraph()  # Spacing

        # Date
        date_str = data.get("date", datetime.now().strftime("%B %d, %Y"))
        self._add_paragraph(doc, date_str, size=11)

        doc.add_paragraph()  # Spacing

        # Recipient Info
        if data.get("hiring_manager"):
            self._add_paragraph(doc, data["hiring_manager"], size=11)
        self._add_paragraph(
            doc, data.get("company", "Company Name"), bold=True, size=11
        )

        doc.add_paragraph()  # Spacing

        # Salutation
        salutation = (
            f"Dear {data.get('hiring_manager', 'Hiring Manager')},"
            if data.get("hiring_manager")
            else "Dear Hiring Manager,"
        )
        self._add_paragraph(doc, salutation, size=11)

        doc.add_paragraph()  # Spacing

        # Letter Content
        content = data.get("content", "")

        # Split content into paragraphs
        paragraphs = content.split("\n\n") if "\n\n" in content else [content]

        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(para_text.strip())
                run.font.name = self.default_font
                run.font.size = Pt(11)
                para.paragraph_format.space_after = Pt(12)

        doc.add_paragraph()  # Spacing

        # Closing
        self._add_paragraph(doc, "Sincerely,", size=11)
        doc.add_paragraph()
        doc.add_paragraph()
        self._add_paragraph(doc, data.get("name", "Your Name"), bold=True, size=11)

        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def generate_proposal(self, data: Dict[str, Any]) -> io.BytesIO:
        """
        Generate business proposal document

        Args:
            data: Proposal data containing:
                - title: Proposal title
                - client_name: Client name
                - prepared_by: Your name/company
                - date: Proposal date
                - content: Proposal content (can be structured or plain text)
                - project_overview: Project description
                - scope: Scope of work
                - deliverables: List of deliverables
                - timeline: Project timeline
                - budget: Budget information

        Returns:
            BytesIO buffer containing the .docx file
        """
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Title Page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(data.get("title", "Business Proposal"))
        title_run.font.name = self.heading_font
        title_run.font.size = Pt(28)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()
        doc.add_paragraph()

        # Prepared For
        for_para = doc.add_paragraph()
        for_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for_run = for_para.add_run(
            f"Prepared for:\n{data.get('client_name', 'Client Name')}"
        )
        for_run.font.name = self.default_font
        for_run.font.size = Pt(14)

        doc.add_paragraph()

        # Prepared By
        by_para = doc.add_paragraph()
        by_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        by_run = by_para.add_run(
            f"Prepared by:\n{data.get('prepared_by', 'Your Company')}"
        )
        by_run.font.name = self.default_font
        by_run.font.size = Pt(14)

        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(
            data.get("date", datetime.now().strftime("%B %d, %Y"))
        )
        date_run.font.name = self.default_font
        date_run.font.size = Pt(12)

        # Page Break
        doc.add_page_break()

        # If structured content is provided
        if data.get("content") and isinstance(data["content"], str):
            # Parse and format the content
            sections_text = data["content"].split("\n\n")
            for section in sections_text:
                if section.strip():
                    lines = section.strip().split("\n")
                    # First line as heading if it looks like a heading
                    if len(lines[0]) < 50 and not lines[0].endswith("."):
                        self._add_heading(doc, lines[0], level=1)
                        for line in lines[1:]:
                            if line.strip():
                                self._add_paragraph(doc, line.strip())
                    else:
                        for line in lines:
                            if line.strip():
                                self._add_paragraph(doc, line.strip())
                    doc.add_paragraph()
        else:
            # Manual structure
            # Executive Summary
            if data.get("executive_summary"):
                self._add_heading(doc, "Executive Summary", level=1)
                self._add_paragraph(doc, data["executive_summary"])
                doc.add_paragraph()

            # Project Overview
            if data.get("project_overview"):
                self._add_heading(doc, "Project Overview", level=1)
                self._add_paragraph(doc, data["project_overview"])
                doc.add_paragraph()

            # Scope of Work
            if data.get("scope"):
                self._add_heading(doc, "Scope of Work", level=1)
                self._add_paragraph(doc, data["scope"])
                doc.add_paragraph()

            # Deliverables
            if data.get("deliverables"):
                self._add_heading(doc, "Deliverables", level=1)
                for deliverable in data["deliverables"]:
                    doc.add_paragraph(deliverable, style="List Bullet")
                doc.add_paragraph()

            # Timeline
            if data.get("timeline"):
                self._add_heading(doc, "Timeline", level=1)
                self._add_paragraph(doc, data["timeline"])
                doc.add_paragraph()

            # Budget
            if data.get("budget"):
                self._add_heading(doc, "Investment", level=1)
                self._add_paragraph(doc, data["budget"])
                doc.add_paragraph()

            # Next Steps
            self._add_heading(doc, "Next Steps", level=1)
            next_steps = data.get(
                "next_steps",
                [
                    "Review this proposal",
                    "Schedule a meeting to discuss details",
                    "Sign agreement and begin work",
                ],
            )
            for step in next_steps:
                doc.add_paragraph(step, style="List Number")

        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def generate_invoice(self, data: Dict[str, Any]) -> io.BytesIO:
        """
        Generate invoice document

        Args:
            data: Invoice data containing:
                - invoice_number: Invoice number
                - invoice_date: Invoice date
                - due_date: Payment due date
                - from_info: Your business info (name, address, email, phone)
                - to_info: Client info (name, address, email)
                - items: List of line items (description, quantity, rate, amount)
                - notes: Additional notes (optional)
                - tax_rate: Tax percentage (optional)

        Returns:
            BytesIO buffer containing the .docx file
        """
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("INVOICE")
        title_run.font.name = self.heading_font
        title_run.font.size = Pt(24)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        # Invoice Info Table
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = "Light Grid Accent 1"

        # From Info (Left)
        from_cell = info_table.cell(0, 0)
        from_info = data.get("from_info", {})
        from_text = "FROM:\n"
        from_text += f"{from_info.get('name', 'Your Business')}\n"
        if from_info.get("address"):
            from_text += f"{from_info['address']}\n"
        if from_info.get("email"):
            from_text += f"{from_info['email']}\n"
        if from_info.get("phone"):
            from_text += f"{from_info['phone']}"
        from_cell.text = from_text

        # Invoice Details (Right)
        details_cell = info_table.cell(0, 1)
        details_text = f"Invoice #: {data.get('invoice_number', 'INV-001')}\n"
        details_text += (
            f"Date: {data.get('invoice_date', datetime.now().strftime('%Y-%m-%d'))}\n"
        )
        details_text += f"Due Date: {data.get('due_date', 'Upon Receipt')}"
        details_cell.text = details_text

        # Bill To (Left)
        to_cell = info_table.cell(1, 0)
        to_info = data.get("to_info", {})
        to_text = "BILL TO:\n"
        to_text += f"{to_info.get('name', 'Client Name')}\n"
        if to_info.get("address"):
            to_text += f"{to_info['address']}\n"
        if to_info.get("email"):
            to_text += f"{to_info['email']}"
        to_cell.text = to_text

        doc.add_paragraph()
        doc.add_paragraph()

        # Items Table
        items = data.get("items", [])
        if items:
            items_table = doc.add_table(rows=len(items) + 1, cols=4)
            items_table.style = "Light Grid Accent 1"

            # Header
            header_cells = items_table.rows[0].cells
            header_cells[0].text = "Description"
            header_cells[1].text = "Quantity"
            header_cells[2].text = "Rate"
            header_cells[3].text = "Amount"

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Items
            subtotal = 0
            for idx, item in enumerate(items, 1):
                row_cells = items_table.rows[idx].cells
                row_cells[0].text = item.get("description", "")
                row_cells[1].text = str(item.get("quantity", 1))
                row_cells[2].text = f"${item.get('rate', 0):.2f}"

                amount = item.get(
                    "amount", item.get("quantity", 1) * item.get("rate", 0)
                )
                row_cells[3].text = f"${amount:.2f}"
                subtotal += amount

            doc.add_paragraph()

            # Totals
            totals_table = doc.add_table(rows=4, cols=2)
            totals_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Subtotal
            totals_table.cell(0, 0).text = "Subtotal:"
            totals_table.cell(0, 1).text = f"${subtotal:.2f}"

            # Tax
            tax_rate = data.get("tax_rate", 0)
            tax_amount = subtotal * (tax_rate / 100) if tax_rate > 0 else 0
            totals_table.cell(1, 0).text = f"Tax ({tax_rate}%):"
            totals_table.cell(1, 1).text = f"${tax_amount:.2f}"

            # Discount
            discount = data.get("discount", 0)
            totals_table.cell(2, 0).text = "Discount:"
            totals_table.cell(2, 1).text = f"-${discount:.2f}"

            # Total
            total = subtotal + tax_amount - discount
            totals_table.cell(3, 0).text = "TOTAL:"
            totals_table.cell(3, 1).text = f"${total:.2f}"

            # Make total row bold
            for cell in [totals_table.cell(3, 0), totals_table.cell(3, 1)]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(14)

        doc.add_paragraph()
        doc.add_paragraph()

        # Notes
        if data.get("notes"):
            self._add_heading(doc, "Notes:", level=2)
            self._add_paragraph(doc, data["notes"])

        # Payment Instructions
        if data.get("payment_instructions"):
            doc.add_paragraph()
            self._add_heading(doc, "Payment Instructions:", level=2)
            self._add_paragraph(doc, data["payment_instructions"])

        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def generate_contract(self, data: Dict[str, Any]) -> io.BytesIO:
        """
        Generate contract document

        Args:
            data: Contract data containing:
                - contract_type: Type of contract
                - date: Contract date
                - party1: First party info (name, address)
                - party2: Second party info (name, address)
                - terms: Contract terms/content
                - effective_date: When contract takes effect
                - expiration_date: When contract expires (optional)

        Returns:
            BytesIO buffer containing the .docx file
        """
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Title
        contract_type = data.get("contract_type", "Service Agreement")
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(contract_type.upper())
        title_run.font.name = self.heading_font
        title_run.font.size = Pt(18)
        title_run.bold = True

        doc.add_paragraph()

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(
            f"Date: {data.get('date', datetime.now().strftime('%B %d, %Y'))}"
        )
        date_run.font.name = self.default_font
        date_run.font.size = Pt(11)

        doc.add_paragraph()

        # Parties
        self._add_heading(doc, "PARTIES", level=1)

        party1 = data.get("party1", {})
        party2 = data.get("party2", {})

        parties_text = f"This Agreement is entered into between:\n\n"
        parties_text += f'Party 1 ("Provider"): {party1.get("name", "Party 1 Name")}'
        if party1.get("address"):
            parties_text += f"\nAddress: {party1['address']}"

        parties_text += f"\n\nAND\n\n"
        parties_text += f'Party 2 ("Client"): {party2.get("name", "Party 2 Name")}'
        if party2.get("address"):
            parties_text += f"\nAddress: {party2['address']}"

        self._add_paragraph(doc, parties_text)
        doc.add_paragraph()

        # Effective Date
        if data.get("effective_date"):
            effective_para = doc.add_paragraph()
            effective_run = effective_para.add_run(
                f"Effective Date: {data['effective_date']}"
            )
            effective_run.font.bold = True
            doc.add_paragraph()

        # Terms
        terms_content = data.get("terms", "")

        if terms_content:
            # Parse terms into sections
            sections_text = terms_content.split("\n\n")
            for section in sections_text:
                if section.strip():
                    lines = section.strip().split("\n")
                    # Check if first line is a heading
                    if len(lines[0]) < 60 and (
                        lines[0].endswith(":") or not lines[0].endswith(".")
                    ):
                        self._add_heading(
                            doc, lines[0].replace(":", "").strip(), level=2
                        )
                        for line in lines[1:]:
                            if line.strip():
                                self._add_paragraph(doc, line.strip())
                    else:
                        for line in lines:
                            if line.strip():
                                self._add_paragraph(doc, line.strip())
                    doc.add_paragraph()

        # Disclaimer
        doc.add_page_break()
        self._add_heading(doc, "LEGAL DISCLAIMER", level=1)
        disclaimer = (
            "This document is provided as a template only and should be reviewed by a "
            "qualified legal professional before use. The parties acknowledge that this "
            "agreement may not be suitable for all situations and that legal advice "
            "should be sought for specific circumstances."
        )
        self._add_paragraph(doc, disclaimer, italic=True, size=10)

        doc.add_paragraph()
        doc.add_paragraph()

        # Signature Section
        self._add_heading(doc, "SIGNATURES", level=1)

        # Party 1 Signature
        doc.add_paragraph()
        self._add_paragraph(doc, "Party 1 (Provider):", bold=True)
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        self._add_paragraph(doc, f"Signature: {party1.get('name', '')}")
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        self._add_paragraph(doc, "Date:")

        doc.add_paragraph()
        doc.add_paragraph()

        # Party 2 Signature
        self._add_paragraph(doc, "Party 2 (Client):", bold=True)
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        self._add_paragraph(doc, f"Signature: {party2.get('name', '')}")
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        self._add_paragraph(doc, "Date:")

        # Save to BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer


# Singleton instance
_docx_generator = None


def get_docx_generator() -> DocxGenerator:
    """Get or create DocxGenerator singleton"""
    global _docx_generator
    if _docx_generator is None:
        _docx_generator = DocxGenerator()
    return _docx_generator
